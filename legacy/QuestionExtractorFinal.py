import json
import pdfplumber
from dotenv import load_dotenv
from langchain.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
import os
from docx import Document
from docx.shared import Inches

current_dir = os.path.dirname(os.path.abspath(__file__))
output_file = os.path.join(current_dir, "output", "output2.json")
output_docx_file = os.path.join(current_dir, "output", "outputW.docx")
pdf_directory = os.path.join(current_dir, "pdfs")
image_dir = os.path.join(current_dir, "images")

# Load environment variables from .env
load_dotenv()

# Create a ChatOpenAI model
model = ChatOpenAI(model="gpt-4o-mini")

# Define the output structure
class ExtractedData(BaseModel):
    filename: str = Field(description="The filename of the processed document")
    question: str = Field(description="The extracted question")
    allocated_points: int = Field(description="The number of points allocated to the question")
    options: list[str] = Field(description="All the options given to select the answer")
    correct_answer: list[int] = Field(description="The correct answers")
    justification: list[str] = Field(description="The justification for the correct answer and why other options are incorrect")
    images: list[str] = Field(description="Paths to images extracted from the document")

# Create a parser
parser = PydanticOutputParser(pydantic_object=ExtractedData)

# Define prompt templates
# prompt_template = ChatPromptTemplate.from_messages([
#     ("system", "You are an expert extractor and a rephraser. Extract the question, allocated points, options, correct answer/s and justifications from the given text. Then rephrase the question, options and justifications while preserving all the information. Format your response as JSON."),
#     ("human", "Use the following text to extract and rephrase the required information and return in json format:\n\n{text}\n\n{format_instructions}")
# ])

prompt_template = ChatPromptTemplate.from_messages([
    ("system", "You are an expert extractor. Extract the question, allocated points, options, correct answer/s and justifications from the given text without changing anything, return the extracted text as it is preserving every word under the extracted field. Format your response as JSON."),
    ("human", "Use the following text to extract and rephrase the required information and return in json format:\n\n{text}\n\n{format_instructions}")
])

table_recognizer_prompt_template = ChatPromptTemplate.from_messages([
    ("system", "You are an expert table recognizer, you need to retrun 'yes' if theres a table in the text, 'no' otherwise."),
    ("human", "Use the following text:\n\n{text}")
])

# Create the LLMChain
chain = (
    prompt_template | model | parser
)

tabel_chain = (
    table_recognizer_prompt_template | model
)


# Function to append JSON to file
def append_json_to_file(json_object, filename=output_file):
    try:
        with open(filename, 'r+') as file:
            # Load existing data
            file_data = json.load(file)
            # Append new data
            file_data.append(json_object)
            # Set file's current position at offset
            file.seek(0)
            # Convert back to JSON and write
            json.dump(file_data, file, indent=4)
    except FileNotFoundError:
        # If file doesn't exist, create it and add the first item
        with open(filename, 'w') as file:
            json.dump([json_object], file, indent=4)


def extract_pdfs(directory, output_file):
    """Converts PDFs in a directory to a single JSON file.

    Args:
        directory: The path to the directory containing PDFs.
        output_file: The path to the output JSON file.
    """

    for filename in os.listdir(directory):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(directory, filename)
            text = ""
            image_paths = []
            print(f"Processing {filename}...")
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Find the position of "NEXT" and extract text up to that point
                        pos = page_text.find("NEXT")
                        if pos != -1:
                            text += page_text[:pos]
                            break
                        text += page_text
                    
                    # Extract images from the page
                    for img_num, img in enumerate(page.images):
                        x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
                        cropped_image = page.within_bbox((x0, y0, x1, y1)).to_image()
                        img_filename = f"{filename}_image_page{page_num+1}_img{img_num+1}.png"
                        img_path = os.path.join(image_dir, img_filename)
                        cropped_image.save(img_path)
                        image_paths.append(img_path)

                    # Run the chain
                result = chain.invoke({
                    "text": text,
                    "format_instructions": parser.get_format_instructions()
                })
                
                # Add filename to the result
                result_dict = result.dict()
                result_dict['filename'] = filename
                result_dict['images'] = image_paths if image_paths else []

                # Append to file
                append_content_to_docx(result_dict, image_paths, output_docx_file)
                append_json_to_file(result_dict)

    # Create the output directory if it doesn't exist
    os.makedirs(os.path.dirname(output_file), exist_ok=True)


def append_content_to_docx(data, image_paths, docx_path):
    # Check if the file exists
    if os.path.exists(docx_path):
        # Open the existing document
        doc = Document(docx_path)
    else:
        # Create a new document
        doc = Document()

    doc.add_heading(data.get('filename', ''), level=1)
    
    # Add the extracted question
    question = data.get('question', '')
    if question:
        doc.add_heading('Question:', level=2)
        doc.add_paragraph(question)
    
    # Add the allocated points
    allocated_points = data.get('allocated_points', None)
    if allocated_points is not None:
        doc.add_heading('Allocated Points:', level=2)
        doc.add_paragraph(str(allocated_points))
    
    # Add the options
    options = data.get('options', [])
    if options:
        doc.add_heading('Options:', level=2)
        for idx, option in enumerate(options, start=1):
            doc.add_paragraph(f"{idx}. {option}")
    
    # Add the correct answers
    correct_answers = data.get('correct_answer', [])
    if correct_answers:
        doc.add_heading('Correct Answers:', level=2)
        doc.add_paragraph(', '.join(str(ans) for ans in correct_answers))
    
    # Add the justifications
    justifications = data.get('justification', [])
    if justifications:
        doc.add_heading('Justifications:', level=2)
        for idx, justification in enumerate(justifications, start=1):
            doc.add_paragraph(f"{idx}. {justification}")
    
    # Add the images
    image_paths = data.get('images', [])
    if image_paths:
        doc.add_heading('Images:', level=2)
        for img_path in image_paths:
            doc.add_picture(img_path)

    # Save the document
    doc.save(docx_path)

extract_pdfs(pdf_directory, output_file)