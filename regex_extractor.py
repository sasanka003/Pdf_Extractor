import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import re
import json
from dotenv import load_dotenv
from langchain.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
import os


current_dir = os.path.dirname(os.path.abspath(__file__))
output_json = os.path.join(current_dir, "output", "output.json")
output_json_with_tables = os.path.join(current_dir, "output", "output_with_tables.json")
output_docx = os.path.join(current_dir, "output", "output.docx")
output_docx_with_tables = os.path.join(current_dir, "output", "output_with_tables.docx")
pdf_directory = os.path.join(current_dir, "pdfs")
image_dir = os.path.join(current_dir, "images")

# Load environment variables from .env
load_dotenv()

# Create a ChatOpenAI model
model = ChatOpenAI(model="gpt-4o-mini")

# Define the output structure
class ExtractedData(BaseModel):
    filename: str = Field(description="The filename of the processed document")
    question: str = Field(description="The extracted question")
    allocated_points: int = Field(description="The number of points allocated to the question")
    options: list[str] = Field(description="All the options given to select the answer")
    correct_answers: list[int] = Field(description="The correct answers")
    justification: list[str] = Field(description="The justification for the correct answer and why other options are incorrect")
    images: list[str] = Field(description="Paths to images extracted from the document")

# Create a parser
parser = PydanticOutputParser(pydantic_object=ExtractedData)

rephraser_prompt_template = ChatPromptTemplate.from_messages([
    ("system", "You are an expert Rephraser. Your task is to rephrase and rewrite the given text preserving all the original data. Ensure that the rephrased text is clear, concise, and maintains the original meaning. If there are python list structures leave them as they are."),
    ("human", "Use the following text:\n\n{text}")
])

table_recognizer_prompt_template = ChatPromptTemplate.from_messages([
    ("system", """
        You are a highly skilled table recognition expert. Your task is to accurately identify and format any tables present in the text. 

        **Key guidelines:**

        * **Precision is paramount:** Ensure that the extracted tables are structurally correct and reflect the original data with utmost fidelity.
        * **Clarity and consistency:** Use a lists of lists style for the tables.
        * **Contextual understanding:** Consider the surrounding text to infer the table's purpose and structure.

        **Output format:**

        * Replace the original table with a formatted list of lists, separated by two newlines.
        * If a table is detected, return the output as a dictionary with two fields: 
          - `content`: The text with the table(s) formatted as list of lists.
          - `table_detected`: `True`
        * If no table is detected, return the output as a dictionary with two fields:
          - `content`: The original text without any modifications.
          - `table_detected`: `False`

        **Example:**

        Original text:

        Name Age City
        Alice 25 New York
        Bob 30 London

        Formatted output:

        [
        ["Name", "Age", "City"],
        ["Alice", "25", "New York"],
        ["Bob", "30", "London"]
        ]
    """),
    ("human", "Use the following text:\n\n{text}\n\n reply in json format")
])

# Create the chain
rephrase_chain = (
    rephraser_prompt_template | model
)

tabel_chain = (
    table_recognizer_prompt_template | model
)



def remove_consecutive_duplicates(text):
    def replace_func(match):
        return match.group(1)
    
    pattern = r'\b(\w+)\s+\1\b'
    return re.sub(pattern, replace_func, text, flags=re.IGNORECASE)

def remove_web_links(text):
    # Remove web links
    text = re.sub(r'http\S+|www\S+', '', text)
    
    return text.strip()

def remove_content_above_question(text):
    # Remove all parts above and including the specific phrase
    phrase = "Answered Review question Quiz-summary"
    text = re.sub(rf'.*?{re.escape(phrase)}', '', text, flags=re.DOTALL)

    return text.strip()

def extract_question_data(text):
    points_pattern = r"(\d+)\s+point\(s\)"

    # Find the points
    match_points = re.search(points_pattern, text)
    # If a match is found, extract the points
    if match_points:
        allocated_points = match_points.group(1)
        text = re.sub(points_pattern, '', text).strip()

    # Regex to capture the question after the question number (e.g., "18. Question")
    question_pattern = r"\d+\.\s*Question\s*(.*?)(?=\n\s*\d+\.\s)"
    question_match = re.search(question_pattern, text, re.DOTALL)
    
    question = ""
    if question_match:
        question = question_match.group(1).strip()
        text = text[question_match.end():]  # Remove the extracted question from the text

    # Regex to find the correct answer
    correct_answer_pattern = r"(\d\.\s.*?)(?=\s✔|\s)"
    correct_answers = re.findall(correct_answer_pattern, text, re.DOTALL)

    # Regex to capture the options
    options_pattern = r"(\d\.\s.*?)(?=\n\s*\d\.|\n\s*(?:CORRECT|INCORRECT))"

    all_options = re.findall(options_pattern, text, re.DOTALL)

    # Create a dictionary for choices
    options = {}

    for option in all_options:
        # Capture the choice number and the choice text
        match = re.match(r"(\d+)\.\s*(.*)", option)
        if match:
            choice_number = match.group(1)
            choice_text = match.group(2)
            # Remove special characters like ✔ and 
            choice_text = re.sub(r"[✔]", "", choice_text).strip()
            # Add the choice to the dictionary
            options[choice_number] = choice_text

    if all_options:
        last_option_match = re.search(all_options[-1], text, re.DOTALL)
        if last_option_match:
            text = text[last_option_match.end():]

    
    # Regex to capture the justification/explanation
    justification_pattern = r"(CORRECT|INCORRECT)\s?(.*)"
    justification_match = re.search(justification_pattern, text, re.DOTALL)
    
    justification_text = ""
    if justification_match:
        justification_text = justification_match.group(2).strip()
        text = text[justification_match.end():]  # Remove the extracted justification from the text

    correct_answer_match = re.search(r'The correct answer is (\d(?: & \d)*)', justification_text)
    if correct_answer_match:
        correct_answer = correct_answer_match.group(1)
    else:
        correct_answer = None

    # Prepare to collect all justifications
    formatted_justifications = {}

    # Pattern to match "Choice X" or "Choice X & Y"
    pattern = re.compile(r"\(Choice[s]? (\d(?: & \d)*)\) (.*?)(?=\(Choice|\Z)", re.DOTALL)

    # Insert the correct choice justification
    correct_answer_pattern = re.compile(r'The correct answer is \d\.\s*(.*?)\s*(?=\(Choice|\Z)', re.DOTALL)
    correct_justification_match = correct_answer_pattern.search(justification_text)
    
    if correct_justification_match and correct_answer:
        correct_justification = correct_justification_match.group(1).strip().replace('\n', ' ')
        formatted_justifications[f"Choice {correct_answer}"]= correct_justification

    # Extract and format other justifications
    for match in pattern.finditer(justification_text):
        choices = match.group(1)
        justification = match.group(2).strip().replace('\n', ' ')
        if choices == correct_answer:
            continue
        formatted_justifications[f"Choice {choices}"]= justification


    data = {
    "question": question,
    "options": options,
    "allocated_points": allocated_points,
    "correct_answers": correct_answers,
    "justifications": formatted_justifications
    }

    # If you want to return the JSON object
    return data

def list_tables_and_rephrase(text, tabel_chain, rephrase_chain):
    # Step 1: Invoke the table chain with the input text
    response = tabel_chain.invoke({
        "text": text,
    })

    print(response.content)
    
    # Step 2: Parse the JSON response
    json_content = response.content.strip('```json').strip('```')
    parsed_data = json.loads(json_content)

    # Extract Tables and content
    tables_present = parsed_data['table_detected']

    # Step 3: Rephrase the content
    rephrased_response = rephrase_chain.invoke({
        "text": parsed_data['content']
    })
    
    # Step 4: Return the rephrased content and tables
    return rephrased_response.content , tables_present

# Function to append JSON to file
def append_json_to_file(json_object, filename):
    try:
        with open(filename, 'r+') as file:
            # Load existing data
            file_data = json.load(file)
            # Append new data
            file_data.append(json_object)
            # Set file's current position at offset
            file.seek(0)
            # Convert back to JSON and write
            json.dump(file_data, file, indent=4)
    except FileNotFoundError:
        # If file doesn't exist, create it and add the first item
        with open(filename, 'w') as file:
            json.dump([json_object], file, indent=4)

def append_content_to_docx(data, image_paths, docx_path):
    # Check if the file exists
    if os.path.exists(docx_path):
        # Open the existing document
        doc = Document(docx_path)
    else:
        # Create a new document
        doc = Document()

    doc.add_heading(data.get('filename', ''), level=1)
    
    # Add the extracted question
    question = data.get('question', '')
    if question:
        doc.add_heading('Question:', level=2)
        doc.add_paragraph(str(question))
    
    # Add the allocated points
    allocated_points = data.get('allocated_points', None)
    if allocated_points is not None:
        doc.add_heading('Allocated Points:', level=2)
        doc.add_paragraph(str(allocated_points))
    
    # Add the options
    options = data.get('options', [])
    if options:
        doc.add_heading('Options:', level=2)
        for idx, option in options.items():
            doc.add_paragraph(f"{idx} - {option}")
    
    # Add the correct answers
    correct_answers = data.get('correct_answer', [])
    if correct_answers:
        doc.add_heading('Correct Answers:', level=2)
        doc.add_paragraph(', '.join(str(ans) for ans in correct_answers))
    
    # Add the justifications
    justifications = data.get('justifications', [])
    if justifications:
        doc.add_heading('Justifications:', level=2)
        for choice, justification in justifications.items():
            doc.add_paragraph(f"{choice} - {justification}")
    
    # Add the images
    image_paths = data.get('images', [])
    if image_paths:
        doc.add_heading('Images:', level=2)
        for img_path in image_paths:
            doc.add_picture(img_path)

    # Save the document
    doc.save(docx_path)

def extract_pdfs(directory):
    """Converts PDFs in a directory to a single JSON file.

    Args:
        directory: The path to the directory containing PDFs.
        output_file: The path to the output JSON file.
    """
        
    for filename in os.listdir(directory):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(directory, filename)
            text = ""
            image_paths = []
            print(f"Processing {filename}...")
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Find the position of "NEXT" and extract text up to that point
                        pos = page_text.find("NEXT")
                        if pos != -1:
                            text += page_text[:pos]
                            break
                        text += page_text
                    
                    # Extract images from the page
                    for img_num, img in enumerate(page.images):
                        x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
                        cropped_image = page.within_bbox((x0, y0, x1, y1)).to_image()
                        img_filename = f"{filename}_image_page{page_num+1}_img{img_num+1}.png"
                        img_path = os.path.join(image_dir, img_filename)
                        cropped_image.save(img_path)
                        image_paths.append(img_path)

            without_top_content = remove_content_above_question(text)
            # Remove consecutive duplicate words
            no_duplicates_text = remove_consecutive_duplicates(without_top_content)

            # Remove web links
            final_text = remove_web_links(no_duplicates_text)
            # print(final_text)

            input_text = extract_question_data(final_text)

            j_consist_tables = False
            q_consist_tables = False

            question, q_consist_tables =list_tables_and_rephrase(input_text["question"], tabel_chain, rephrase_chain)
            
            result_dict = input_text
            result_dict['question'] = question

            for key in input_text['justifications']:
                justification, consist_tables =list_tables_and_rephrase(input_text['justifications'][key], tabel_chain, rephrase_chain)
                result_dict['justifications'][key] = justification
                if consist_tables:
                    j_consist_tables = True
            
            result_dict['filename'] = filename
            result_dict['images'] = image_paths if image_paths else []

            if q_consist_tables or j_consist_tables:
                result_dict['consist_tables'] = True
                append_content_to_docx(result_dict, image_paths, output_docx_with_tables)
                append_json_to_file(result_dict, output_json_with_tables)
            else:
                result_dict['consist_tables'] = False
                append_content_to_docx(result_dict, image_paths, output_docx)
                append_json_to_file(result_dict, output_json)                

            print(result_dict)
                    
    return

extract_pdfs(pdf_directory)