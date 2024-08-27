# import pdfplumber
# from docx import Document
# from docx.shared import Inches
# from PIL import Image
# import io
# import re
# import os


# current_dir = os.path.dirname(os.path.abspath(__file__))
# pdf_file = os.path.join(current_dir, "pdfs", "64.pdf")
# docx_file = os.path.join(current_dir, "output", "output73.docx")

# def extract_content_until_next(pdf_path):
#     text = ""
#     images = []
#     tables = []
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 # Find the position of "NEXT" and extract text up to that point
#                 pos = page_text.find("NEXT")
#                 if pos != -1:
#                     text += page_text[:pos]
#                     break
#                 text += page_text
            
#             # Extract images from the page
#             for img in page.images:
#                 x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
#                 cropped_image = page.within_bbox((x0, y0, x1, y1)).to_image()
#                 img_bytes = io.BytesIO()
#                 cropped_image.save(img_bytes, format='PNG')
#                 img_bytes.seek(0)
#                 images.append(img_bytes)
            
#             # Extract tables from the page
#             page_tables = page.extract_tables()
#             print(page_tables)
#             tables.extend(page_tables)
#     return text, images, tables

# # def filter_text(text):
# #     filtered_lines = []
# #     for line in text.split('\n'):
# #         # Check if the line contains any letters
# #         if re.search('[a-zA-Z]', line):
# #             filtered_lines.append(line)
# #     return '\n'.join(filtered_lines)

# # def remove_consecutive_duplicates(text):
# #     def replace_func(match):
# #         return match.group(1)
    
# #     pattern = r'\b(\w+)\s+\1\b'
# #     return re.sub(pattern, replace_func, text, flags=re.IGNORECASE)

# # def remove_web_links_and_phrases(text):
# #     # Remove web links
# #     text = re.sub(r'http\S+|www\S+', '', text)
# #     # Remove specific phrases
# #     phrases_to_remove = ["Answered Review question Quiz-summary", "1 point(s)"]
# #     for phrase in phrases_to_remove:
# #         text = text.replace(phrase, '')
# #     return text

# def save_content_to_docx(text, images, tables, docx_path):
#     doc = Document()
#     doc.add_paragraph(text)
#     for img_bytes in images:
#         doc.add_picture(img_bytes, width=Inches(6))

#     for table in tables:
#         doc_table = doc.add_table(rows=1, cols=len(table[0]))
#         hdr_cells = doc_table.rows[0].cells
#         for i, cell_text in enumerate(table[0]):
#             hdr_cells[i].text = cell_text

#         for row in table[1:]:
#             row_cells = doc_table.add_row().cells
#             for i, cell_text in enumerate(row):
#                 row_cells[i].text = cell_text

#     doc.save(docx_path)

# # Extract text, images, and tables from PDF until the phrase "NEXT"
# extracted_text, extracted_images, extracted_tables = extract_content_until_next(pdf_file)

# # # Filter out lines that contain only numbers
# # filtered_text = filter_text(extracted_text)

# # # Remove consecutive duplicate words
# # no_duplicates_text = remove_consecutive_duplicates(filtered_text)

# # # Remove web links and specific phrases
# # final_text = remove_web_links_and_phrases(no_duplicates_text)



# # Save the final text, images, and tables to a DOCX file
# # save_content_to_docx(final_text, extracted_images, extracted_tables, docx_file)

# # print(f"Filtered content up to 'NEXT' with specified content removed has been saved to {docx_file}")



import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
import re
import os

current_dir = os.path.dirname(os.path.abspath(__file__))
pdf_file = os.path.join(current_dir, "pdfs", "73.pdf")
docx_file = os.path.join(current_dir, "output", "output73.docx")
image_dir = os.path.join(current_dir, "images")

def extract_content_until_next(pdf_path):
    text = ""
    image_paths = []
    tables = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                # Find the position of "NEXT" and extract text up to that point
                pos = page_text.find("NEXT")
                if pos != -1:
                    text += page_text[:pos]
                    break
                text += page_text
            
            # Extract images from the page
            for img_num, img in enumerate(page.images):
                x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
                cropped_image = page.within_bbox((x0, y0, x1, y1)).to_image()
                img_filename = f"image_page{page_num+1}_img{img_num+1}.png"
                img_path = os.path.join(image_dir, img_filename)
                cropped_image.save(img_path)
                image_paths.append(img_path)
            
            # Extract tables from the page
            page_tables = page.extract_tables()
            tables.extend(page_tables)
    return text, image_paths, tables

def filter_text(text):
    filtered_lines = []
    for line in text.split('\n'):
        # Check if the line contains any letters
        if re.search('[a-zA-Z]', line):
            filtered_lines.append(line)
    return '\n'.join(filtered_lines)

def remove_consecutive_duplicates(text):
    def replace_func(match):
        return match.group(1)
    
    pattern = r'\b(\w+)\s+\1\b'
    return re.sub(pattern, replace_func, text, flags=re.IGNORECASE)

def remove_web_links_and_phrases(text):
    # Remove web links
    text = re.sub(r'http\S+|www\S+', '', text)
    # Remove specific phrases
    phrases_to_remove = ["Answered Review question Quiz-summary", "1 point(s)"]
    for phrase in phrases_to_remove:
        text = text.replace(phrase, '')
    return text

def save_content_to_docx(text, image_paths, tables, docx_path):
    doc = Document()
    doc.add_paragraph(text)
    for img_path in image_paths:
        doc.add_picture(img_path, width=Inches(6))

    for table in tables:
        doc_table = doc.add_table(rows=1, cols=len(table[0]))
        hdr_cells = doc_table.rows[0].cells
        for i, cell_text in enumerate(table[0]):
            hdr_cells[i].text = cell_text

        for row in table[1:]:
            row_cells = doc_table.add_row().cells
            for i, cell_text in enumerate(row):
                row_cells[i].text = cell_text

    doc.save(docx_path)

# Create the image directory if it doesn't exist
os.makedirs(image_dir, exist_ok=True)

# Extract text, images, and tables from PDF until the phrase "NEXT"
extracted_text, extracted_image_paths, extracted_tables = extract_content_until_next(pdf_file)

# Filter out lines that contain only numbers
filtered_text = filter_text(extracted_text)

# Remove consecutive duplicate words
no_duplicates_text = remove_consecutive_duplicates(filtered_text)

# Remove web links and specific phrases
final_text = remove_web_links_and_phrases(no_duplicates_text)

# Save the final text, images, and tables to a DOCX file
save_content_to_docx(final_text, extracted_image_paths, extracted_tables, docx_file)

print(f"Filtered content up to 'NEXT' with specified content removed has been saved to {docx_file}")
print(f"Images have been saved to {image_dir}")