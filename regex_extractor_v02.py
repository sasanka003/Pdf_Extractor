import pdfplumber
from docx import Document
import re
import json
from dotenv import load_dotenv
from langchain.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
import os
import firebase_admin
from firebase_admin import credentials, firestore, storage
from bson import ObjectId
from datetime import datetime
import time
from pymongo import MongoClient


current_dir = os.path.dirname(os.path.abspath(__file__))
# output_json = os.path.join(current_dir, "output", "output.json")
# output_json_with_tables = os.path.join(current_dir, "output", "output_with_tables.json")
output_docx = os.path.join(current_dir, "output", "output.docx")
output_docx_with_tables = os.path.join(current_dir, "output", "output_with_tables.docx")
pdf_directory = os.path.join(current_dir, "pdfs", "with table")
# pdf_directory = os.path.join(current_dir, "pdfs")
image_dir = os.path.join(current_dir, "images")
firebase_credentials = os.path.join(current_dir, "firebase_credentials.json")

# Load environment variables from .env
load_dotenv()

FIRESTORE_BUCKET = os.getenv('FIRESTORE_BUCKET')
MONGO_URI = os.getenv('MONGO_URI')

BANK_NAME = "CUSTOM"

# Initialize Firebase
cred = credentials.Certificate(firebase_credentials)
firebase_admin.initialize_app(cred,
                              {
        'storageBucket': FIRESTORE_BUCKET
    })

# Initialize Storage
bucket = storage.bucket()

# MongoDB setup
client = MongoClient(MONGO_URI)
db = client['amc-site2']  # Database name
collection = db['questions']  # Collection name

# Create a ChatOpenAI model
model = ChatOpenAI(model="gpt-4o")

rephraser_prompt_template = ChatPromptTemplate.from_messages([
    ("system", """
        You are an expert Rephraser, skilled in linguistic clarity and precision. Your task is to rephrase the provided text while ensuring that all factual data, context, and meaning are strictly preserved. Avoid adding or omitting any details.

        **Guidelines:**
        - **Clarity:** Ensure the rephrased text is clear and easy to understand.
        - **Conciseness:** Reduce unnecessary words or redundancy without altering the meaning.
        - **Preserve Structure:** Leave markdown elements, such as tables, lists, or special formats, unchanged and rephrase only the prose sections.
        - **Maintain Meaning:** Be especially careful that the rephrased version does not change the intent or tone of the original text.

        If the input includes tables in markdown format, keep them intact and focus on rephrasing the non-table parts of the text.
    """),
    ("human", "Use the following text:\n\n{text}")
])


table_recognizer_prompt_template = ChatPromptTemplate.from_messages([
    ("system", """
        You are a table recognition and formatting specialist. Your task is to detect, extract, and properly format any tables found in the provided text, ensuring the highest level of accuracy.

        **Key Instructions:**
        - **Accuracy First:** Ensure all tables are correctly identified and represented in markdown format without any data loss.
        - **No Alterations:** The text surrounding the table should remain unchanged unless explicitly related to the table content.
        - **Table Structure:** Understand the context to ensure that the table structure (rows and columns) is logical and reflects the data's meaning.
        - **Consistency:** Maintain consistency across all tables, using markdown formatting.
        
        **Output Specification:**
        - If one or more tables are detected:
          - Return a JSON object with:
            - `content`: The text with tables formatted in markdown, followed by two newlines.
            - `table_detected`: `true`
        - If no table is found:
          - Return a JSON object with:
            - `content`: The original text without changes.
            - `table_detected`: `false`

        **Note:** Use only 'true' or 'false' for the 'table_detected' field. Always provide the result in JSON format.
        
        **Example of Table Formatting:**
        Original Text:
        
        Name Age City
        Alice 25 New York
        Bob 30 London
        
        Markdown Format:
        
        | Name  | Age | City      |
        |-------|-----|-----------|
        | Alice | 25  | New York  |
        | Bob   | 30  | London    |
    """),
    ("human", "Use the following text:\n\n{text}")
])

# Create the chain
rephrase_chain = (
    rephraser_prompt_template | model
)

tabel_chain = (
    table_recognizer_prompt_template | model
)

# Function to generate ObjectId
def generate_object_id():
    return ObjectId()

# Function to get the current timestamp
def get_current_timestamp():
    return datetime.utcnow()

def remove_consecutive_duplicates(text):
    def replace_func(match):
        return match.group(1)
    
    pattern = r'\b(\w+)\s+\1\b'
    return re.sub(pattern, replace_func, text, flags=re.IGNORECASE)

def remove_web_links(text):
    # Remove web links
    text = re.sub(r'http\S+|www\S+', '', text)
    
    return text.strip()

def remove_content_above_question(text):
    # Remove all parts above and including the specific phrase
    phrase = "Answered Review question Quiz-summary"
    text = re.sub(rf'.*?{re.escape(phrase)}', '', text, flags=re.DOTALL)

    return text.strip()

def extract_question_data(text):
    points_pattern = r"(\d+)\s+point\(s\)"
    
    # Find the points (Note: Removed points from final output, but kept extraction for future use)
    match_points = re.search(points_pattern, text)
    if match_points:
        allocated_points = match_points.group(1)
        text = re.sub(points_pattern, '', text).strip()

    # Regex to capture the question after the question number (e.g., "18. Question")
    question_pattern = r"\d+\.\s*Question\s*(.*?)(?=\n\s*\d+\.\s)"
    question_match = re.search(question_pattern, text, re.DOTALL)
    
    question_text = ""
    if question_match:
        question_text = question_match.group(1).strip()
        text = text[question_match.end():]  # Remove the extracted question from the text

    # Regex to capture the options (answers)
    options_pattern = r"(\d\.\s.*?)(?=\n\s*\d\.|\n\s*(?:CORRECT|INCORRECT))"
    all_options = re.findall(options_pattern, text, re.DOTALL)

    # Prepare a list for storing answers in the required format
    answers = []

    for option in all_options:
        # Capture the choice number and the choice text
        match = re.match(r"(\d+)\.\s*(.*)", option)
        if match:
            choice_number = match.group(1)
            choice_text = match.group(2)
            # Remove special characters like ✔ and 
            choice_text = re.sub(r"[✔]", "", choice_text).strip()
            choice_text  = re.sub(r"\s*[\uF00D]", "", choice_text).strip()
            # Add the choice to the answers list with a placeholder for isCorrect and explanation
            answers.append({
                "text": f"{choice_number} {choice_text}",  # Add the choice text
                "isCorrect": False,  # This will be updated later
                "explanation": ""    # Will add justifications later
            })

    if all_options:
        last_option_match = re.search(all_options[-1], text, re.DOTALL)
        if last_option_match:
            text = text[last_option_match.end():]

    # Regex to capture the justification/explanation
    justification_pattern = r"(CORRECT|INCORRECT)\s?(.*)"
    justification_match = re.search(justification_pattern, text, re.DOTALL)
    
    justification_text = ""
    if justification_match:
        justification_text = justification_match.group(2).strip()
        text = text[justification_match.end():]  # Remove the extracted justification from the text

    # print(justification_text)

    # Extract the correct answer number
    correct_answer_match = re.search(r'The correct answer is (\d(?: & \d)*)', justification_text)
    if correct_answer_match:
        correct_answer = correct_answer_match.group(1)
    else:
        correct_answer = None

    # print("-------------")
    # print(correct_answer)

    # Collect justifications
    formatted_justifications = {}

    # Pattern to match "Choice X" or "Choice X & Y"
    pattern = re.compile(r"\(Choice[s]? (\d(?: & \d)*)\) (.*?)(?=\(Choice|\Z)", re.DOTALL)

    # Insert the correct choice justification
    correct_choice_pattern = re.compile(r'The correct answer is \d\.\s*(.*?)\s*(?=\(Choice|\Z)', re.DOTALL)
    correct_justification_match = correct_choice_pattern.search(justification_text)
    
    if correct_justification_match and correct_answer:
        correct_justification = correct_justification_match.group(1).strip().replace('\n', ' ')
        formatted_justifications[f"Choice {correct_answer}"]= correct_justification

    # Extract and format other justifications
    for match in pattern.finditer(justification_text):
        choices = match.group(1)
        justification = match.group(2).strip().replace('\n', ' ')
        if "&" in choices or "," in choices:
            # Handle multiple choices case, replacing "&" and "and" with commas and splitting by ","
            choices = choices.replace("&", ",").replace("and", ",").strip()
            
            # Loop through each choice and assign the same justification
            for choice in choices.split(','):
                choice = choice.strip()
                formatted_justifications[f"Choice {choice}"] = justification
        else:
            # General case for a single choice
            formatted_justifications[f"Choice {choices}"] = justification

    # print("-------------")
    # print(json.dumps(formatted_justifications, indent=4))

    # Now update the answers list with the correct answers and justifications
    for answer in answers:
        # print("-------------")
        # print(answer)
        # Extract choice number safely by checking if it starts with a digit
        match = re.search(r"(\d+)", answer["text"])
        if match:
            choice_number = match.group(1)  # Extract choice number
            # print("-------------")
            # print(choice_number)
            answer["isCorrect"] = choice_number == correct_answer  # Mark correct answer
            answer["explanation"] = formatted_justifications.get(f"Choice {choice_number}", "")  # Add explanation if exists
            answer["text"] = re.sub(r"^\d+\s+", "", answer["text"])

    # Return the data in the required format
    data = {
        "question": question_text,
        "answers": [
            {
                "text": answer["text"],
                "isCorrect": answer["isCorrect"],
                "explanation": answer["explanation"],
            }
            for answer in answers
        ],
    }

    return data

# Function to append JSON to file
def append_json_to_file(json_object, filename):
    try:
        with open(filename, 'r+') as file:
            # Load existing data
            file_data = json.load(file)
            # Append new data
            file_data.append(json_object)
            # Set file's current position at offset
            file.seek(0)
            # Convert back to JSON and write
            json.dump(file_data, file, indent=4)
    except FileNotFoundError:
        # If file doesn't exist, create it and add the first item
        with open(filename, 'w') as file:
            json.dump([json_object], file, indent=4)

def list_tables_and_rephrase(text, tabel_chain, rephrase_chain, retries=3):
    """Rephrases content and detects tables, with retry mechanism for errors."""
    attempt = 0  # Track the number of retries
    while attempt < retries:
        try:
            # Step 1: Invoke the table chain with the input text
            response = tabel_chain.invoke({
                "text": text,
            })

            print(response.content)

            # Step 2: Parse the JSON response
            json_content = response.content.strip('```json').strip('```')

            try:
                parsed_data = json.loads(json_content)
                print("JSON file parsed successfully.")
                print(parsed_data['table_detected'])
            except json.JSONDecodeError:
                raise ValueError("Invalid JSON format detected.")
            
            # Extract Tables and content
            tables_present = parsed_data['table_detected']

            # Step 3: Rephrase the content
            rephrased_response = rephrase_chain.invoke({
                "text": parsed_data['content']
            })

            # Step 4: Return the rephrased content and tables
            return rephrased_response.content, tables_present, True

        except (json.JSONDecodeError, KeyError, ValueError) as e:
            print(f"Error encountered: {e}. Retrying... {attempt + 1}/{retries}")
            attempt += 1
            time.sleep(2)  # Wait for a short duration before retrying

        except Exception as e:
            print(f"An unexpected error occurred: {e}.")
            break

    print("Max retries reached or an unrecoverable error occurred.")
    return None, None, False

def append_content_to_docx(data, image_paths, docx_path):
    # Check if the file exists
    if os.path.exists(docx_path):
        # Open the existing document
        doc = Document(docx_path)
    else:
        # Create a new document
        doc = Document()

    # Add the filename as a heading
    doc.add_heading(data.get('filename', ''), level=1)
    
    # Add the extracted question
    question_text = data.get('question', '')
    if question_text:
        doc.add_heading('Question:', level=2)
        doc.add_paragraph(question_text)
    
    # Add the answers (options with their correctness and explanations)
    answers = data.get('answers', [])
    if answers:
        doc.add_heading('Answers:', level=2)
        for answer in answers:
            text = answer.get('text', '')
            is_correct = answer.get('isCorrect', False)
            explanation = answer.get('explanation', '')

            # Mark if the answer is correct
            correctness_label = "(Correct)" if is_correct else "(Incorrect)"
            doc.add_paragraph(f"{text} {correctness_label}")

            # Add explanation if available
            if explanation:
                doc.add_paragraph(f"Explanation: {explanation}", style='Quote')
    
    # Add the images
    if image_paths:
        doc.add_heading('Images:', level=2)
        for img_path in image_paths:
            doc.add_picture(img_path)

    # Save the document
    doc.save(docx_path)


def upload_image_to_firebase(image_path, filename, question_id):
    """Uploads an image to Firebase Storage and returns the image URL."""
    
    # Define the folder path in Firebase Storage
    folder_path = f"question-assets/{BANK_NAME}/{question_id}/"
    
    # Create the full path by appending the folder path and filename
    firebase_storage_path = f"{folder_path}{filename}"
    
    # Create the blob (the file object) with the full path
    blob = bucket.blob(firebase_storage_path)
    
    # Upload the image from the local image path
    blob.upload_from_filename(image_path)
    
    # Get the public URL and Firebase path
    firebase_url = blob.public_url  # The URL to access the image publicly
    firebase_path = blob.name       # Path in the Firebase bucket
    
    return {
        "url": firebase_url,
        "path": firebase_path  # The full path inside the Firebase bucket
    }

incomplete = []

def extract_pdfs(directory):
    """Converts PDFs in a directory to a single JSON file.

    Args:
        directory: The path to the directory containing PDFs.
        output_file: The path to the output JSON file.
    """
        
    for filename in os.listdir(directory):
        if filename.endswith(".pdf"):
            result_dict = {}
            result_dict["_id"] = generate_object_id()
            pdf_path = os.path.join(directory, filename)
            text = ""
            image_paths = []
            assets = []  # To store the asset objects
            print(f"Processing {filename}...")
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Find the position of "NEXT" and extract text up to that point
                        pos = page_text.find("NEXT")
                        if pos != -1:
                            text += page_text[:pos]
                            break
                        text += page_text

                    for img_num, img in enumerate(page.images):
                        x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
                        cropped_image = page.within_bbox((x0, y0, x1, y1)).to_image()
                        
                        # Generate filename and save path
                        img_filename = f"{filename}_image_page{page_num+1}_img{img_num+1}.png"
                        img_path = os.path.join(image_dir, img_filename)
                        
                        # Save the cropped image locally
                        cropped_image.save(img_path)
                        image_paths.append(img_path)
                        
                        # Upload the image and get the URL and path in the Firebase bucket
                        upload_result = upload_image_to_firebase(img_path, img_filename, result_dict["_id"])
                        image_url = upload_result["url"]
                        firebase_path = upload_result["path"]  # This is the path inside the Firebase bucket
                        
                        # Get the image size in bytes
                        image_size = os.path.getsize(img_path)
                        image_id = generate_object_id()
                        
                        # Add the image data into the assets list
                        assets.append({
                            "url": image_url,
                            "path": firebase_path,  # Firebase bucket path
                            "size": image_size,
                            "_id": image_id
                        })

            without_top_content = remove_content_above_question(text)
            # Remove consecutive duplicate words
            no_duplicates_text = remove_consecutive_duplicates(without_top_content)

            # Remove web links
            final_text = remove_web_links(no_duplicates_text)
            # print(final_text)

            input_text = extract_question_data(final_text)
            print("--------------------------input_text---------------------------")
            print(input_text)
            print("---------------------------------------------------------------")

            j_consist_tables = False
            q_consist_tables = False
            complete = True

            question, q_consist_tables, complete =list_tables_and_rephrase(input_text["question"], tabel_chain, rephrase_chain)
            
            result_dict['filename'] = filename
            result_dict['question'] = question

            for answer in input_text['answers']:
                if 'explanation' in answer and answer['explanation']:
                    # Rephrase the explanation and check for tables
                    explanation, consist_tables, complete = list_tables_and_rephrase(answer['explanation'], tabel_chain, rephrase_chain)
                    
                    # Update the explanation in the answer
                    answer['explanation'] = explanation
                    answer['_id'] = generate_object_id()
                    
                    # Set flag if tables are found in the explanation
                    if consist_tables:
                        j_consist_tables = True
        
            result_dict['answers'] = input_text['answers']
            result_dict['assets'] = assets

            result_dict["bank"] = "CUSTOM"
            result_dict["status"] = "AVAILABLE"
            result_dict["addedBy"] = "SCRIPT"
            result_dict["createdAt"] = get_current_timestamp()
            result_dict["updatedAt"] = get_current_timestamp()

            if(complete):
                if q_consist_tables or j_consist_tables:
                    result_dict['consist_tables'] = True
                    append_content_to_docx(result_dict, image_paths, output_docx_with_tables)
                    # append_json_to_file(result_dict, output_json_with_tables)
                    collection.insert_one(result_dict)
                else:
                    result_dict['consist_tables'] = False
                    append_content_to_docx(result_dict, image_paths, output_docx)
                    # append_json_to_file(result_dict, output_json)
                    collection.insert_one(result_dict)
            else:
                incomplete.append(filename)

            print("--------------------------output_text---------------------------")
            print(result_dict)
            print("----------------------------------------------------------------")
        print("Incomplete files: ", incomplete)    

    return

extract_pdfs(pdf_directory)